from docx import Document

doc = Document()
doc.add_heading('LLM Engineering & Hugging Face Cheat Sheet', 0)

sections = {
    "1. Model Types": [
        "CausalLM: Chatbots, text generation (generate, TextStreamer, chat templates)",
        "Seq2Seq: Translation, summarization (generate, pipelines)",
        "Diffusion: Image generation (diffusers pipelines only)",
        "Embedding Models: Semantic search (vector outputs only)"
    ],
    "2. Tokenizers": [
        "tokenizer(): text → tokens (all models)",
        "decode(): tokens → text (all models)",
        "apply_chat_template(): ONLY works for chat/instruction-tuned models with templates"
    ],
    "3. Generation": [
        "generate(): works for CausalLM & Seq2Seq only"
    ],
    "4. Streaming": [
        "TextStreamer: works mainly with CausalLM"
    ],
    "5. Pipelines": [
        "text-generation → CausalLM",
        "summarization → Seq2Seq",
        "translation → Seq2Seq",
        "diffusion pipelines → diffusion models only"
    ],
    "6. Agentic AI": [
        "LLM (usually CausalLM)",
        "Tools",
        "Memory (short-term + long-term)",
        "Planner",
        "Executor"
    ],
    "7. RAG": [
        "Embedding model",
        "Vector DB (FAISS, Pinecone, etc.)",
        "Retriever",
        "Generator (CausalLM or Seq2Seq)"
    ],
    "8. Key Insight": [
        "Model decides capability",
        "Tokenizer decides format",
        "Pipeline decides convenience",
        "Agent system decides intelligence"
    ]
}

for title, points in sections.items():
    doc.add_heading(title, level=1)
    for point in points:
        doc.add_paragraph(point)

doc.save("LLM_CheatSheet.docx")
print("Saved as LLM_CheatSheet.docx")